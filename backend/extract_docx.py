from docx import Document
import sys

def extract_text(file_path):
    try:
        doc = Document(file_path)
        fullText = []
        for para in doc.paragraphs:
            if para.text.strip():
                fullText.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    fullText.append(" | ".join(row_text))
                    
        return "\n".join(fullText)
    except Exception as e:
        return str(e)

if __name__ == "__main__":
    path = "d:/Credit Flow/Appraisal/MTL - Yakub Rain.docx"
    content = extract_text(path)
    with open("extracted_text.txt", "w", encoding="utf-8") as f:
        f.write(content)
    print("Extraction successful - see extracted_text.txt")
